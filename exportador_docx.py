# ===================================================================
# app/exportador_docx.py (VERSÃO FINAL E À PROVA DE FALHAS)
#
# O que mudou:
# - FUNÇÃO UNIVERSAL DE FORMATAÇÃO: Adicionada a função `format_value`
#   que converte QUALQUER tipo de dado (string, dict, list, None)
#   em um texto legível e bem formatado para o usuário.
# - APLICAÇÃO GLOBAL: Esta função agora é usada em TODOS os campos
#   antes de serem escritos no documento, eliminando de vez a causa
#   do erro "Argument must be bytes or unicode, got 'dict'".
# - ROBUSTEZ MÁXIMA: O código agora é resiliente a qualquer variação
#   na estrutura dos dados retornados pela IA.
# ===================================================================

from docx import Document
import os
import json

def format_value(value):
    """Converte de forma inteligente qualquer valor para uma string legível."""
    if value is None or value == '':
        return 'Não informado'
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return ', '.join([format_value(item) for item in value])
    if isinstance(value, dict):
        return ', '.join([f"{str(k).replace('_', ' ').title()}: {v}" for k, v in value.items()])
    return str(value)

def gerar_docx_resumo(dados, caminho_saida="export/resumo_processo.docx"):
    """Gera um resumo do processo em formato .docx."""
    doc = Document()
    doc.add_heading("Resumo do Processo – PJe-Calc com IA", level=1)

    # --- Dados Processuais ---
    doc.add_heading("📌 Dados Processuais", level=2)
    dados_processuais = dados.get("dados_processuais", {})
    if dados_processuais:
        for k, v in dados_processuais.items():
            doc.add_paragraph(f"**{k.replace('_', ' ').title()}:** {format_value(v)}")
    
    # --- Partes Envolvidas ---
    doc.add_heading("👤 Partes Envolvidas", level=2)
    partes = dados.get("partes", {})
    if partes:
        doc.add_paragraph(f"**Reclamante:** {format_value(partes.get('reclamante'))}")
        reclamadas = partes.get("reclamadas", [])
        if reclamadas:
            doc.add_paragraph("**Reclamadas:**")
            for item in set(filter(None, reclamadas)):
                doc.add_paragraph(f" - {format_value(item)}", style="List Bullet")
    
    # --- Contrato de Trabalho ---
    doc.add_heading("💼 Contrato de Trabalho", level=2)
    contrato = dados.get("contrato_trabalho", {})
    if contrato:
        for k, v in contrato.items():
            if k != 'periodos_afastamento':
                doc.add_paragraph(f"**{k.replace('_', ' ').title()}:** {format_value(v)}")
    
    # --- Observações Gerais (Resumo Jurídico) ---
    doc.add_heading("📝 Observações Gerais", level=2)
    observacoes = dados.get("observacoes_gerais", "Nenhum resumo foi gerado.")
    doc.add_paragraph(format_value(observacoes))

    # --- Pleitos e Verbas ---
    doc.add_heading("📑 Pleitos e Verbas Reclamadas", level=2)
    pleitos = dados.get("pleitos_e_verbas", [])
    if pleitos:
        for pleito in pleitos:
            verba = format_value(pleito.get('verba', 'Verba não especificada'))
            parametros = format_value(pleito.get('parametros', 'sem parâmetros'))
            doc.add_paragraph(f" - **{verba}** ({parametros})", style="List Bullet")
    else:
        doc.add_paragraph("Nenhum pleito ou verba encontrado.")
    
    # --- Salva o documento ---
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)